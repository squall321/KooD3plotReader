"""Generate fibonacci_mollweide_analysis.docx with proper equation formatting."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import lxml.etree as ET


def add_equation(doc, equation_xml_list):
    """Add an Office Math equation paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    # Add math namespace
    for eq_xml in equation_xml_list:
        math_elem = parse_xml(eq_xml)
        p._element.append(math_elem)
    return p


def omath(content):
    """Wrap content in oMath element."""
    return f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{content}</m:oMath>'


def run(text, bold=False, italic=True):
    """Create a math run."""
    rPr = ''
    if bold:
        rPr = '<m:rPr><m:sty m:val="bi"/></m:rPr>'
    elif not italic:
        rPr = '<m:rPr><m:sty m:val="p"/></m:rPr>'
    return f'<m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{rPr}<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>{text}</m:t></m:r>'


def frac(num, den):
    """Create a fraction."""
    return f'<m:f xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'


def rad(content):
    """Square root."""
    return f'<m:rad xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:radPr><m:degHide m:val="1"/></m:radPr><m:deg/><m:e>{content}</m:e></m:rad>'


def sup(base, superscript):
    """Superscript."""
    return f'<m:sSup xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:e>{base}</m:e><m:sup>{superscript}</m:sup></m:sSup>'


def sub(base, subscript):
    """Subscript."""
    return f'<m:sSub xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:e>{base}</m:e><m:sub>{subscript}</m:sub></m:sSub>'


def func(name, arg):
    """Function like sin, cos, arccos."""
    return f'<m:func xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:fName>{run(name, italic=False)}</m:fName><m:e>{arg}</m:e></m:func>'


def delim(content, open_char='(', close_char=')'):
    """Delimiters (parentheses, brackets)."""
    return f'<m:d xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:dPr><m:begChr m:val="{open_char}"/><m:endChr m:val="{close_char}"/></m:dPr><m:e>{content}</m:e></m:d>'


def nary(op, sub_content, sup_content, base):
    """Summation, product, integral."""
    return f'<m:nary xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:naryPr><m:chr m:val="{op}"/></m:naryPr><m:sub>{sub_content}</m:sub><m:sup>{sup_content}</m:sup><m:e>{base}</m:e></m:nary>'


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2B579A')

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_shading(cell, 'D9E2F3')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6D)

    # ====== TITLE ======
    title = doc.add_heading('Fibonacci Sphere Sampling + Mollweide Projection을 이용한\n전각도 낙하 응력 분석', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('KooD3plotReader / KooReport Technical Document')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    r.font.italic = True

    doc.add_paragraph()  # spacer

    # ====== 1. 문제 정의 ======
    doc.add_heading('1. 문제 정의: 왜 전각도 낙하 시험이 필요한가', level=1)

    doc.add_paragraph(
        '제품 낙하 시험에서 충격 방향은 단위 구(unit sphere) S² 위의 한 점으로 표현된다. '
        '기존 규격(IEC 60068-2-31 등)은 면(face) 6방향, 모서리(edge) 12방향, 꼭짓점(corner) 8방향 — '
        '총 26방향만 평가한다.'
    )

    # d = (sinθcosφ, sinθsinφ, cosθ) ∈ S²
    add_equation(doc, [omath(
        run('d', bold=True) + run(' = ') +
        delim(
            func('sin', run('θ')) + func('cos', run('φ')) + run(', ') +
            func('sin', run('θ')) + func('sin', run('φ')) + run(', ') +
            func('cos', run('θ'))
        ) + run(' ∈ ') + sup(run('S', italic=False), run('2'))
    )])

    p = doc.add_paragraph()
    p.add_run('여기서 θ는 극각(polar angle), φ는 방위각(azimuthal angle)이다.').italic = False

    doc.add_paragraph(
        '26방향은 정육면체의 대칭축에 해당하며, 이들 사이의 "빈 공간"에 최악 방향이 존재할 수 있다. '
        '예를 들어 면과 모서리 사이의 중간 각도, 또는 꼭짓점에서 약간 벗어난 방향에서 응력이 최대가 되는 '
        '경우가 흔하다.'
    )

    p = doc.add_paragraph()
    r = p.add_run('전각도 낙하 분석의 목적: ')
    r.bold = True
    p.add_run('구 전체를 균일하게 샘플링하여, 방향 의존적 응력 분포의 전체 맵을 구성하고, 최악 방향을 놓치지 않는 것.')

    # ====== 2. Fibonacci Lattice Sampling ======
    doc.add_heading('2. Fibonacci Lattice Sampling', level=1)

    doc.add_heading('2.1 알고리즘', level=2)
    doc.add_paragraph(
        'Fibonacci lattice는 황금비(golden ratio)를 이용하여 구 위에 N개의 점을 '
        '준균일(quasi-uniform)하게 배치하는 방법이다.'
    )

    # φ = (1+√5)/2 ≈ 1.6180
    add_equation(doc, [omath(
        run('ϕ') + run(' = ') +
        frac(run('1 + ') + rad(run('5')), run('2')) +
        run(' ≈ 1.6180    (황금비)', italic=False)
    )])

    doc.add_paragraph('N개의 점의 구면 좌표 (k = 0, 1, …, N−1):')

    # θ_k = arccos(1 - (2k+1)/N)
    add_equation(doc, [omath(
        sub(run('θ'), run('k')) + run(' = ') +
        func('arccos', delim(
            run('1 − ') + frac(run('2k + 1'), run('N'))
        ))
    )])

    # φ_k = 2πk / φ²
    add_equation(doc, [omath(
        sub(run('φ'), run('k')) + run(' = 2πk ⋅ ') +
        frac(run('1'), sup(run('ϕ'), run('2'))) +
        run(' ≈ 2πk ⋅ 0.3820')
    )])

    doc.add_paragraph('여기서:')
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    r = p.add_run('θ')
    r.italic = True
    r = p.add_run('k')
    r.font.subscript = True
    p.add_run(' — 극각. cos θ를 [1, −1] 구간에서 균등 분할하여 ')
    r = p.add_run('위도 방향 균일성')
    r.bold = True
    p.add_run('을 보장')

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    r = p.add_run('φ')
    r.italic = True
    r = p.add_run('k')
    r.font.subscript = True
    p.add_run(' — 방위각. 황금비의 제곱의 역수(ϕ⁻²)만큼 회전하여 ')
    r = p.add_run('경도 방향 비반복 분산')
    r.bold = True
    p.add_run('을 달성')

    doc.add_heading('2.2 왜 황금비인가', level=2)
    doc.add_paragraph(
        '황금비는 "가장 무리수적인 무리수(most irrational number)"로, 연분수 전개에서 모든 계수가 1이다. '
        '이 성질 때문에 연속적인 점들이 가장 균일하게 분산된다. 각 새로운 점은 기존 점들 사이의 '
        '가장 큰 간격(largest gap)을 분할하는 경향이 있어, 임의의 N에서도 준최적 분포를 달성한다.'
    )

    doc.add_heading('2.3 Cuboid 26방향과의 비교', level=2)
    add_styled_table(doc,
        ['항목', 'Cuboid 26', 'Fibonacci 100', 'Fibonacci 1,146'],
        [
            ['방향 수', '26', '100', '1,146'],
            ['평균 이웃 간격', '15.9°', '7.9°', '~6°'],
            ['구 커버리지', '~12%', '~12%', '~99%'],
            ['분포 특성', '정육면체 대칭축만', '준균일', '준균일, 고밀도'],
            ['최악 방향 탐지', '제한적', '양호', '우수'],
            ['시뮬레이션 비용', '낮음', '중간', '높음'],
        ],
        col_widths=[3.5, 3.0, 3.0, 3.0]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Cuboid의 한계: ')
    r.bold = True
    p.add_run(
        '26방향은 정육면체 기하학에 종속되어 있어, 면과 모서리 사이의 약 15°~30° 범위의 방향을 전혀 평가하지 못한다. '
        '제품 형상이 직육면체가 아닌 경우(비대칭 내부 구조, 편심 질량 등) 최악 방향이 이 "사각지대"에 놓일 확률이 높다.'
    )

    p = doc.add_paragraph()
    r = p.add_run('Fibonacci의 장점: ')
    r.bold = True
    p.add_run(
        'N을 임의로 설정 가능하며, 어떤 N에서도 구 위의 점 분포가 준균일하다. 또한 격자 구조가 없으므로 '
        'aliasing(격자 아티팩트)이 없고, 방향 해상도 ↔ 시뮬레이션 비용의 트레이드오프를 자유롭게 조절할 수 있다.'
    )

    doc.add_heading('2.4 최소 샘플 수 결정', level=2)
    doc.add_paragraph('평균 이웃 간격 Δθ는 다음으로 근사된다:')

    # Δθ ≈ √(4π/N)
    add_equation(doc, [omath(
        run('Δθ ≈ ') + rad(frac(run('4π'), run('N'))) +
        run('    (rad)')
    )])

    add_styled_table(doc,
        ['목표 해상도', '필요 N'],
        [
            ['30°', '~46'],
            ['15°', '~183'],
            ['10°', '~413'],
            ['6°', '~1,146'],
            ['3°', '~4,584'],
            ['1°', '~41,253'],
        ],
        col_widths=[4, 4]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        '참고: 구의 총 제곱도(square degrees)는 4π × (180/π)² ≈ 41,253이므로, '
        '1° 해상도는 구 전체를 1° × 1° 셀로 분할하는 것과 동일한 밀도이다.'
    )
    doc.add_paragraph(
        '전각도 낙하 시험에서는 일반적으로 6°~10° 해상도가 산업적으로 충분하며, 이는 N = 400~1,200에 해당한다.'
    )

    # ====== 3. Mollweide Projection ======
    doc.add_heading('3. Mollweide Projection (몰바이데 도법)', level=1)

    doc.add_heading('3.1 정의', level=2)
    doc.add_paragraph(
        'Mollweide 도법은 등적(equal-area) 의사원통 도법으로, 구 위의 면적비가 평면에서도 보존된다. '
        '지도학에서 전세계 분포 패턴 시각화에 사용되며, 본 시스템에서는 충격 방향별 응력 분포를 시각화한다.'
    )

    doc.add_paragraph('경도 λ, 위도 φ에 대한 투영 좌표 (x, y):')

    # x = (2√2/π) λ cos θ
    add_equation(doc, [omath(
        run('x = ') + frac(run('2') + rad(run('2')), run('π')) +
        run(' λ ') + func('cos', run('θ'))
    )])

    # y = √2 sin θ
    add_equation(doc, [omath(
        run('y = ') + rad(run('2')) + run(' ') + func('sin', run('θ'))
    )])

    doc.add_paragraph('여기서 보조각(auxiliary angle) θ는 다음 초월방정식의 해이다:')

    # 2θ + sin(2θ) = π sin φ
    add_equation(doc, [omath(
        run('2θ + ') + func('sin', delim(run('2θ'))) + run(' = π ') + func('sin', run('φ'))
    )])

    doc.add_paragraph('이 방정식은 해석적 해가 없으므로, Newton-Raphson 반복법으로 풀이한다:')

    # θ_{n+1} = θ_n - (2θ_n + sin(2θ_n) - π sin φ) / (2 + 2cos(2θ_n))
    add_equation(doc, [omath(
        sub(run('θ'), run('n+1')) + run(' = ') + sub(run('θ'), run('n')) + run(' − ') +
        frac(
            run('2') + sub(run('θ'), run('n')) + run(' + ') + func('sin', delim(run('2') + sub(run('θ'), run('n')))) + run(' − π ') + func('sin', run('φ')),
            run('2 + 2') + func('cos', delim(run('2') + sub(run('θ'), run('n'))))
        )
    )])

    doc.add_heading('3.2 등적 성질의 의미', level=2)
    doc.add_paragraph('Mollweide 도법의 핵심은 등적(equal-area) 성질이다. 이것이 응력 분석에서 중요한 이유:')

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    r = p.add_run('시각적 면적 비례: ')
    r.bold = True
    p.add_run('맵에서 빨간(고응력) 영역의 넓이가 실제 구 위에서 위험 방향이 차지하는 입체각(solid angle)에 비례한다.')

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    r = p.add_run('정량적 해석 가능: ')
    r.bold = True
    p.add_run('"전체 방향의 30%에서 항복 응력을 초과"와 같은 진술이 맵의 면적 비율로 직접 읽힌다.')

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    r = p.add_run('편향 없는 비교: ')
    r.bold = True
    p.add_run('극지방과 적도 지역의 면적이 동일하게 표현되어, 특정 방향이 과대/과소 대표되지 않는다.')

    doc.add_paragraph(
        '참고로, 일반 정사각형 위경도 격자(equirectangular projection)에서는 극지방이 심하게 왜곡되어 면적 비율이 의미를 잃는다.'
    )

    doc.add_heading('3.3 역변환 (Inverse Mollweide)', level=2)
    doc.add_paragraph('IDW 보간에서 맵 픽셀 → 구면 좌표 역변환이 필요하다:')

    # θ = arcsin(y/√2)
    add_equation(doc, [omath(
        run('θ = ') + func('arcsin', delim(frac(run('y'), rad(run('2')))))
    )])

    # φ = arcsin((2θ + sin(2θ))/π)
    add_equation(doc, [omath(
        run('φ = ') + func('arcsin', delim(
            frac(run('2θ + ') + func('sin', delim(run('2θ'))), run('π'))
        ))
    )])

    # λ = πx / (2√2 cos θ)
    add_equation(doc, [omath(
        run('λ = ') + frac(run('πx'), run('2') + rad(run('2')) + run(' ') + func('cos', run('θ')))
    )])

    doc.add_heading('3.4 Newton-Raphson 수렴 안정화', level=2)
    doc.add_paragraph(
        '표준 Newton-Raphson은 특정 위도에서 발산할 수 있다. '
        "도함수 f'(θ) = 2 + 2cos(2θ)가 θ = ±π/4 부근에서 작아지면 "
        '스텝이 과도하게 커져 θ가 유효 범위 [−π/2, π/2]를 벗어나고, '
        'sin(huge number)가 무작위 값을 반환하여 y좌표가 뒤집히는 현상이 발생한다.'
    )

    p = doc.add_paragraph()
    r = p.add_run('해결: 감쇠 Newton-Raphson (Damped Newton-Raphson)')
    r.bold = True

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('스텝 크기를 최대 0.3 rad로 제한')
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('θ를 항상 [−π/2 + ε, π/2 − ε] 범위로 클램핑')
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('극점(|φ| > π/2 − ε)에서는 직접 θ = ±π/2 할당')

    # ====== 4. IDW 구면 보간 ======
    doc.add_heading('4. IDW (Inverse Distance Weighting) 구면 보간', level=1)

    doc.add_heading('4.1 알고리즘', level=2)
    doc.add_paragraph('N개의 데이터 포인트 (λᵢ, φᵢ, vᵢ)가 주어졌을 때, 임의의 위치 (λ₀, φ₀)에서의 보간값:')

    # v̂ = Σwᵢvᵢ / Σwᵢ
    add_equation(doc, [omath(
        run('v̂') + delim(sub(run('λ'), run('0')) + run(', ') + sub(run('φ'), run('0'))) +
        run(' = ') +
        frac(
            nary('∑', sub(run('i'), run('=1')), run('N'),
                 sub(run('w'), run('i')) + run(' ⋅ ') + sub(run('v'), run('i'))),
            nary('∑', sub(run('i'), run('=1')), run('N'),
                 sub(run('w'), run('i')))
        )
    )])

    # wᵢ = 1/d^p
    add_equation(doc, [omath(
        sub(run('w'), run('i')) + run(' = ') +
        frac(run('1'),
             sup(run('d') + delim(sub(run('p'), run('0')) + run(', ') + sub(run('p'), run('i'))), run('p')))
    )])

    doc.add_paragraph('여기서 d(p₀, pᵢ)는 구면 대원 거리(great-circle distance)로, Haversine 공식으로 계산한다:')

    # Haversine formula
    add_equation(doc, [omath(
        run('d = 2 ') + func('arcsin', rad(
            sup(func('sin', frac(run('Δφ'), run('2'))), run('2')) +
            run(' + ') +
            func('cos', sub(run('φ'), run('1'))) +
            func('cos', sub(run('φ'), run('2'))) +
            sup(func('sin', frac(run('Δλ'), run('2'))), run('2'))
        ))
    )])

    doc.add_heading('4.2 파라미터 선택', level=2)
    add_styled_table(doc,
        ['파라미터', '값', '이유'],
        [
            ['거듭제곱 p', '3.5', '높은 값 → 날카로운 피크, 데이터 포인트에서 정확한 값 보존'],
            ['Snap 반경', '0.02 rad (~1.1°)', '데이터 포인트 근처에서 보간 대신 정확값 사용'],
            ['그리드 스텝', '3 px', '성능과 정밀도의 균형. 이중선형 보간으로 픽셀 해상도 복원'],
        ],
        col_widths=[3, 3, 8]
    )

    doc.add_paragraph()
    doc.add_paragraph('거듭제곱 p의 선택은 시각화 품질에 직접 영향을 미친다:')

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    r = p.add_run('p = 2 (표준): ')
    r.bold = True
    p.add_run('부드러운 보간이지만 피크가 퍼져서 최대값이 보존되지 않음')

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    r = p.add_run('p = 3.5 (본 시스템): ')
    r.bold = True
    p.add_run('데이터 포인트 위치에서 원래 값이 거의 정확히 재현되면서도 사이 영역은 자연스럽게 보간')

    doc.add_heading('4.3 왜 구면 거리인가', level=2)
    doc.add_paragraph(
        'Mollweide 맵 좌표 상의 유클리드 거리 대신 구면 대원 거리를 사용하는 이유:'
    )
    doc.add_paragraph(
        'Mollweide 도법은 등적(equal-area)이지만 등거(equidistant)는 아니다. '
        '즉, 맵 상의 직선 거리가 실제 구면 거리와 일치하지 않는다. 유클리드 거리를 사용하면:'
    )
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('극지방 근처에서 가까운 점이 멀게, 먼 점이 가깝게 계산됨')
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('경도 ±180° 경계에서 불연속 발생')
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('보간 결과에 투영 왜곡이 그대로 반영')

    doc.add_paragraph('구면 거리를 사용하면 투영 방식에 무관하게 물리적으로 정확한 보간이 이루어진다.')

    # ====== 5. 충격 방향 매핑 ======
    doc.add_heading('5. 충격 방향 매핑', level=1)

    doc.add_heading('5.1 3D 방향 벡터 → (경도, 위도)', level=2)
    doc.add_paragraph('충격 방향 벡터 d = (dₓ, dᵧ, d_z)를 몰바이데 맵 좌표로 변환:')

    # λ = atan2(dₓ, -d_z)
    add_equation(doc, [omath(
        run('λ = ') + func('atan2', delim(sub(run('d'), run('x')) + run(', −') + sub(run('d'), run('z'))))
    )])

    # φ = arcsin(dᵧ)
    add_equation(doc, [omath(
        run('φ = ') + func('arcsin', delim(sub(run('d'), run('y'))))
    )])

    p = doc.add_paragraph()
    r = p.add_run('좌표 컨벤션:')
    r.bold = True

    add_styled_table(doc,
        ['맵 좌표', '물리적 방향', '설명'],
        [
            ['중심 (0°, 0°)', 'Back (−Z)', '맵 중심 = 뒷면 충격'],
            ['동쪽 (+90°)', 'Right (+X)', '오른쪽 충격'],
            ['서쪽 (−90°)', 'Left (−X)', '왼쪽 충격'],
            ['북극 (+90°)', 'Top (+Y)', '상면 충격'],
            ['남극 (−90°)', 'Bottom (−Y)', '하면 충격'],
            ['±180°', 'Front (+Z)', '전면 충격'],
        ],
        col_widths=[3.5, 3.0, 4.0]
    )

    doc.add_heading('5.2 Euler 각도에서의 변환', level=2)
    doc.add_paragraph('시뮬레이션에서 충격 방향은 Euler 각도 (roll, pitch, yaw)로 정의된다. 초기 충격 벡터 [0, 0, −1]에 회전 행렬을 적용하여 전개하면:')

    add_equation(doc, [omath(
        sub(run('d'), run('x')) + run(' = −') + func('sin', delim(run('pitch', italic=False)))
    )])
    add_equation(doc, [omath(
        sub(run('d'), run('y')) + run(' = ') + func('sin', delim(run('roll', italic=False))) + func('cos', delim(run('pitch', italic=False)))
    )])
    add_equation(doc, [omath(
        sub(run('d'), run('z')) + run(' = −') + func('cos', delim(run('roll', italic=False))) + func('cos', delim(run('pitch', italic=False)))
    )])

    # ====== 6. 파이프라인 요약 ======
    doc.add_heading('6. 분석 파이프라인 요약', level=1)

    add_styled_table(doc,
        ['단계', '입력', '출력', '핵심 알고리즘'],
        [
            ['1. 방향 생성', 'N (샘플 수)', 'N개 (θ, φ) 좌표', 'Fibonacci lattice'],
            ['2. 시뮬레이션', '방향별 d3plot', '응력/변형률/가속도 시계열', 'LS-DYNA explicit solver'],
            ['3. 후처리', 'analysis_result.json', '부품별 피크값', 'Unified Analyzer'],
            ['4. 좌표 변환', 'Euler (roll, pitch, yaw)', '(경도, 위도)', 'Rotation matrix → atan2/arcsin'],
            ['5. 투영', '(경도, 위도, 값)', 'Mollweide (x, y, 색상)', 'Equal-area projection + IDW'],
            ['6. 시각화', 'Mollweide contour map', 'HTML 리포트', 'Canvas + SVG overlay'],
        ],
        col_widths=[2.5, 3.5, 3.5, 4.0]
    )

    # ====== 7. 결론 ======
    doc.add_heading('7. 결론', level=1)
    doc.add_paragraph(
        'Fibonacci lattice sampling과 Mollweide equal-area projection의 조합은 '
        '전각도 낙하 시뮬레이션 분석에 다음과 같은 이점을 제공한다:'
    )

    items = [
        ('균일 샘플링', '구 위의 모든 방향을 편향 없이 평가하여 최악 방향을 놓치지 않음'),
        ('확장 가능성', '26개(cuboid) → 100개 → 1,000개로 점진적으로 해상도를 높일 수 있음'),
        ('등적 시각화', '맵에서 위험 영역의 넓이가 실제 입체각에 비례하여 정량적 해석 가능'),
        ('투영 무관 보간', '구면 대원 거리 기반 IDW가 투영 왜곡을 완전히 회피'),
        ('산업 적용성', '기존 26방향 규격과 호환되면서도, 제품 형상 의존적 취약 방향을 추가로 탐지'),
    ]
    for i, (title_text, desc) in enumerate(items, 1):
        p = doc.add_paragraph()
        p.style = 'List Number'
        r = p.add_run(f'{title_text}: ')
        r.bold = True
        p.add_run(desc)

    # ====== References ======
    doc.add_heading('참고 문헌', level=1)
    refs = [
        'González, Á. (2010). "Measurement of areas on a sphere using Fibonacci and latitude–longitude lattices." Mathematical Geosciences, 42(1), 49–64.',
        'Swinbank, R., & Purser, R. J. (2006). "Fibonacci grids: A novel approach to global modelling." Quarterly Journal of the Royal Meteorological Society, 132(619), 1769–1793.',
        'Snyder, J. P. (1987). Map Projections — A Working Manual. U.S. Geological Survey Professional Paper 1395. (Mollweide projection: pp. 249–252)',
        'Shepard, D. (1968). "A two-dimensional interpolation function for irregularly-spaced data." Proceedings of the 1968 ACM National Conference, 517–524.',
        'IEC 60068-2-31:2008. Environmental testing — Part 2-31: Tests — Test Ec: Rough handling shocks, primarily for equipment-type specimens.',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f'[{i}] {ref}')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        for r in p.runs:
            r.font.size = Pt(9)

    return doc


if __name__ == '__main__':
    out_path = Path(__file__).parent / 'fibonacci_mollweide_analysis.docx'
    doc = build_document()
    doc.save(str(out_path))
    print(f'Saved: {out_path}')
